"""
Generate Word documents for validation materials.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Load data
sample = json.load(open('data/analysis/rigorous_capacity/validation_sample.json', encoding='utf-8'))
corpus = json.load(open('data/corpus/corpus_master_20260128.json', encoding='utf-8'))
corpus_lookup = {e['url']: e for e in corpus['entries']}
enriched = json.load(open('data/oecd/enriched/oecd_enriched_20260127_203406.json', encoding='utf-8'))
enriched_lookup = {p['title'] + '_' + p['jurisdiction']: p for p in enriched['policies']}

output_dir = Path('data/analysis/validation')

# =============================================================================
# DOCUMENT 1: Coding Protocol
# =============================================================================
print("Creating coding protocol...")
doc = Document()

# Title
title = doc.add_heading('Implementation Capacity Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Manual Validation Protocol', level=1)
doc.add_paragraph()

# Purpose
doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'This document provides instructions for manually coding a sample of 50 AI governance '
    'policies to validate our automated implementation capacity scores. Your independent '
    'assessment will help us measure the reliability of our methodology.'
)

# Instructions
doc.add_heading('Instructions for Coders', level=2)
doc.add_paragraph('For each policy in the "Policy Texts" document:', style='List Number')
doc.add_paragraph('Read the complete policy text carefully', style='List Number 2')
doc.add_paragraph('Score each of the 5 dimensions using the 0-4 scale below', style='List Number 2')
doc.add_paragraph('Record evidence (quotes or line numbers) supporting your score', style='List Number 2')
doc.add_paragraph('Enter your scores in the provided spreadsheet', style='List Number 2')
doc.add_paragraph('Flag any uncertainties with a "?" in the notes column', style='List Number 2')

doc.add_paragraph()
doc.add_paragraph(
    'Important: Do NOT look at the automated scores until AFTER you complete your manual coding. '
    'The automated scores are hidden in the policy document - only reveal them for comparison after finishing.',
    style='Intense Quote'
)

# Coding scheme
doc.add_heading('Coding Scheme', level=2)

dimensions = [
    {
        'name': 'CLARITY',
        'question': 'Does the policy clearly specify what it aims to achieve?',
        'look_for': 'Objectives, goals, targets, timelines, definitions, scope statements',
        'scores': [
            ('0', 'No clear objectives stated'),
            ('1', 'General objectives without specifics (e.g., "Promote AI development")'),
            ('2', 'Specific objectives but no measurable targets'),
            ('3', 'Measurable targets for some objectives (e.g., "Train 10,000 specialists by 2025")'),
            ('4', 'Comprehensive targets with timelines for all major objectives'),
        ]
    },
    {
        'name': 'RESOURCES',
        'question': 'Does the policy specify what resources will be allocated?',
        'look_for': 'Budget amounts, funding sources, staff/FTE numbers, infrastructure, equipment',
        'scores': [
            ('0', 'No resources mentioned'),
            ('1', 'General statement about need for resources'),
            ('2', 'Commitment to allocate without specifics'),
            ('3', 'Specific amounts for some resource types (e.g., "€50 million for research")'),
            ('4', 'Comprehensive allocation: budget, staff, infrastructure with sources'),
        ]
    },
    {
        'name': 'AUTHORITY',
        'question': 'Does the policy specify who is responsible and what powers they have?',
        'look_for': 'Agency names, ministries, commissions, enforcement powers, penalties, legal basis',
        'scores': [
            ('0', 'No authority structures mentioned'),
            ('1', 'General reference to government responsibility'),
            ('2', 'Named agency without specific powers'),
            ('3', 'Named agency with some defined powers (e.g., "may issue guidance")'),
            ('4', 'Clear authority with enforcement powers and sanctions'),
        ]
    },
    {
        'name': 'ACCOUNTABILITY',
        'question': 'Does the policy specify how implementation will be monitored?',
        'look_for': 'Monitoring, evaluation, reporting requirements, review cycles, KPIs, audits',
        'scores': [
            ('0', 'No accountability mechanisms'),
            ('1', 'General commitment to monitoring'),
            ('2', 'Monitoring mentioned without specifics'),
            ('3', 'Specific monitoring with some reporting (e.g., "annual report")'),
            ('4', 'Comprehensive M&E: KPIs, review cycles, evaluation methodology'),
        ]
    },
    {
        'name': 'COHERENCE',
        'question': 'Is the policy aligned with other policies and international standards?',
        'look_for': 'References to other laws, coordination mechanisms, international standards (ISO, OECD, EU)',
        'scores': [
            ('0', 'Isolated policy with no references'),
            ('1', 'Mentions other policies without integration'),
            ('2', 'Some coordination mechanisms mentioned'),
            ('3', 'Explicit alignment with specific policies'),
            ('4', 'Comprehensive coherence: cross-references, coordination body, international alignment'),
        ]
    },
]

for dim in dimensions:
    doc.add_heading(f'Dimension: {dim["name"]}', level=3)
    doc.add_paragraph(dim['question'], style='Intense Quote')
    doc.add_paragraph(f'Look for: {dim["look_for"]}')
    
    # Create scoring table
    table = doc.add_table(rows=len(dim['scores']) + 1, cols=2)
    table.style = 'Table Grid'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Score'
    header_cells[1].text = 'Criteria'
    
    # Data rows
    for i, (score, criteria) in enumerate(dim['scores'], 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = score
        row_cells[1].text = criteria
    
    doc.add_paragraph()

# Workflow
doc.add_heading('Workflow', level=2)
doc.add_paragraph('1. Coder A and Coder B each code all 50 policies independently')
doc.add_paragraph('2. Do NOT discuss policies during coding')
doc.add_paragraph('3. Submit completed spreadsheets to lead researcher')
doc.add_paragraph('4. Inter-rater reliability will be calculated')
doc.add_paragraph('5. Discrepancies (≥2 point difference) will be discussed')
doc.add_paragraph('6. Consensus scores will be compared to automated scores')

# Time estimate
doc.add_heading('Time Estimate', level=2)
doc.add_paragraph('Expect approximately 5-8 minutes per policy, or 4-6 hours total.')

doc.save(output_dir / 'Validation_Protocol.docx')
print(f"  Saved: {output_dir / 'Validation_Protocol.docx'}")


# =============================================================================
# DOCUMENT 2: Policy Texts for Coding
# =============================================================================
print("\nCreating policy texts document...")
doc2 = Document()

# Title
title = doc2.add_heading('Policy Texts for Validation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc2.add_paragraph('50 AI Governance Policies for Manual Coding')
doc2.add_paragraph()
doc2.add_paragraph(
    'Instructions: Read each policy text and score using the coding scheme in the Protocol document. '
    'Enter your scores in the Excel spreadsheet. Do NOT reveal the automated scores until after '
    'completing your manual assessment.',
    style='Intense Quote'
)
doc2.add_paragraph()

# Add each policy
for i, p in enumerate(sample, 1):
    doc2.add_heading(f'P{i:03d}: {p["title"]}', level=1)
    
    # Metadata table
    table = doc2.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Jurisdiction'
    table.rows[0].cells[1].text = p['jurisdiction']
    table.rows[1].cells[0].text = 'Year'
    table.rows[1].cells[1].text = str(p.get('year', 'Unknown'))
    table.rows[2].cells[0].text = 'Word Count'
    table.rows[2].cells[1].text = str(p.get('word_count', 0))
    table.rows[3].cells[0].text = 'Policy ID'
    table.rows[3].cells[1].text = f'P{i:03d}'
    
    doc2.add_paragraph()
    
    # Get full text
    key = p['title'] + '_' + p['jurisdiction']
    enriched_policy = enriched_lookup.get(key, {})
    url = enriched_policy.get('url', '')
    corpus_entry = corpus_lookup.get(url, {})
    
    full_text = corpus_entry.get('content', '') or enriched_policy.get('initiative_overview', '') or ''
    
    doc2.add_heading('Full Text', level=2)
    
    if full_text:
        # Limit to ~2500 words for readability
        words = full_text.split()
        if len(words) > 2500:
            doc2.add_paragraph(f'[Showing first 2500 of {len(words)} words]', style='Caption')
            full_text = ' '.join(words[:2500]) + '...'
        
        # Split into paragraphs for better formatting
        paragraphs = full_text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc2.add_paragraph(para.strip())
    else:
        doc2.add_paragraph('[No text available in corpus]', style='Caption')
    
    # Scoring box
    doc2.add_heading('Your Scores', level=2)
    score_table = doc2.add_table(rows=6, cols=3)
    score_table.style = 'Table Grid'
    
    headers = score_table.rows[0].cells
    headers[0].text = 'Dimension'
    headers[1].text = 'Score (0-4)'
    headers[2].text = 'Evidence'
    
    dims = ['Clarity', 'Resources', 'Authority', 'Accountability', 'Coherence']
    for j, dim in enumerate(dims, 1):
        score_table.rows[j].cells[0].text = dim
        score_table.rows[j].cells[1].text = ''
        score_table.rows[j].cells[2].text = ''
    
    doc2.add_paragraph()
    
    # Hidden automated scores (for comparison after)
    doc2.add_heading('Automated Scores (REVEAL AFTER MANUAL CODING)', level=3)
    auto_total = p['clarity_score'] + p['resources_score'] + p['authority_score'] + p['accountability_score'] + p['coherence_score']
    
    auto_table = doc2.add_table(rows=6, cols=2)
    auto_table.style = 'Light Shading'
    auto_table.rows[0].cells[0].text = 'Clarity'
    auto_table.rows[0].cells[1].text = f'{p["clarity_score"]}/4'
    auto_table.rows[1].cells[0].text = 'Resources'
    auto_table.rows[1].cells[1].text = f'{p["resources_score"]}/4'
    auto_table.rows[2].cells[0].text = 'Authority'
    auto_table.rows[2].cells[1].text = f'{p["authority_score"]}/4'
    auto_table.rows[3].cells[0].text = 'Accountability'
    auto_table.rows[3].cells[1].text = f'{p["accountability_score"]}/4'
    auto_table.rows[4].cells[0].text = 'Coherence'
    auto_table.rows[4].cells[1].text = f'{p["coherence_score"]}/4'
    auto_table.rows[5].cells[0].text = 'TOTAL'
    auto_table.rows[5].cells[1].text = f'{auto_total}/20'
    
    doc2.add_page_break()

doc2.save(output_dir / 'Policy_Texts_for_Coding.docx')
print(f"  Saved: {output_dir / 'Policy_Texts_for_Coding.docx'}")


# =============================================================================
# DOCUMENT 3: Create Excel spreadsheet
# =============================================================================
print("\nCreating Excel spreadsheet...")
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Validation Scores"
    
    # Headers
    headers = [
        'Policy_ID', 'Title', 'Jurisdiction', 'Year',
        'Clarity (0-4)', 'Resources (0-4)', 'Authority (0-4)', 
        'Accountability (0-4)', 'Coherence (0-4)', 'TOTAL',
        'Notes'
    ]
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color='DAEEF3', end_color='DAEEF3', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
    
    # Data rows
    for i, p in enumerate(sample, 1):
        row = i + 1
        ws.cell(row=row, column=1, value=f'P{i:03d}')
        ws.cell(row=row, column=2, value=p['title'][:60])
        ws.cell(row=row, column=3, value=p['jurisdiction'])
        ws.cell(row=row, column=4, value=p.get('year', ''))
        # Leave score columns empty for manual entry
        ws.cell(row=row, column=10, value='=SUM(E{0}:I{0})'.format(row))
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 14
    ws.column_dimensions['G'].width = 14
    ws.column_dimensions['H'].width = 16
    ws.column_dimensions['I'].width = 14
    ws.column_dimensions['J'].width = 10
    ws.column_dimensions['K'].width = 30
    
    wb.save(output_dir / 'Validation_Scoring_Sheet.xlsx')
    print(f"  Saved: {output_dir / 'Validation_Scoring_Sheet.xlsx'}")
    
except ImportError:
    print("  (openpyxl not installed - skipping Excel file)")

print("\n" + "=" * 60)
print("VALIDATION PACKAGE COMPLETE")
print("=" * 60)
print(f"\nFiles in {output_dir}:")
print("  1. Validation_Protocol.docx - Instructions for coders")
print("  2. Policy_Texts_for_Coding.docx - 50 policies with text")
print("  3. Validation_Scoring_Sheet.xlsx - Spreadsheet for scores")
print("\nDistribute all 3 files to each coder.")
